import json
import re
import ssl
from urllib.request import Request, urlopen
from bs4 import BeautifulSoup
from docx import Document
import time

BASE_URL="https://fr.indeed.com/"

ctx=ssl.create_default_context()
ctx.check_hostname = False
ctx.verify_mode = ssl.CERT_NONE


def main():
    name_of_city = input("Entrez la ville où vous cherchez un travail")
    keywords = input("Entrez les mots-clé")
    number_of_pages = input("Le nombre de pages que vous souhaitez scraper")
    name_of_city = '+'.join(name_of_city.split(' '))
    keywords = '+'.join(keywords.split(' '))
    url_to_scrape = BASE_URL + "/jobs?q=" + keywords + "&l=" + name_of_city
    number_of_pages_nos = int(number_of_pages)
    data_collected = scrape_data(url_to_scrape, number_of_pages_nos)
    print(len(data_collected))

    for i in range(0, len(data_collected)):
        title = data_collected[i]['title']
        title_for_doc = data_collected[i]['title']
        title = title.replace(" ", "_")
        title = title.replace("/", "-")
        details = data_collected[i]['details']

        document = Document()
        document.add_heading(title_for_doc, 0)
        document.add_paragraph(details)
        document.save(title + '.docx')

        # with open(title, "w") as f:
        # f.write(details)
        # f.close()
    with open('data.json', "w") as fp:
        json.dump(data_collected, fp, sort_keys=True, indent=4, ensure_ascii=False)

def scrape_data(url_to_scrape,number_of_pages_nos):
    data_collected = []
    for i in range(0,number_of_pages_nos):
        extension = ""
        if i is not 0:
            extension = "&start-"+str(i*10)
        url = url_to_scrape + extension
        req = Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        web_page = urlopen(req).read()
        soup = BeautifulSoup(web_page, 'html.parser')
        data_collected = get_data_from_webpage(data_collected, soup)
        time.sleep(1)
    return data_collected


def extract_data_points(job, div):
    for a in div.findAll('a', attrs={'class': 'jobtitle turnstileLink'}):
        job['title'] = a['title']
    for a1 in div.findAll('a', attrs={'data-tn-element': 'companyName'}):
        job['companyName'] = a1.text.strip()
    for span in div.findAll('span', attrs={'class': 'ratingsContent'}):
        job['rating'] = span.text.strip()
    for span1 in div.findAll('span', attrs={'class': 'location accessible-contrast-color-location'}):
        job['location'] = span1.text.strip()
    for div in div.findAll('div', attrs={'class': 'jobsearch-jobDescriptionText'}):
        summary = div.text.strip()
        job['summary'] = re.sub(' +', ' ', summary.replace("\n", "\n"))
    for span2 in div.findAll('span', attrs={'class': 'date'}):
        job['date'] = span2.text.strip()

    return job


def get_data_from_webpage(data_collected, soup):
    job_posts = []
    for div in soup.findAll('div', attrs={'class': 'jobsearch-SerpJobCard unifiedRow row result'}):
        job = dict()
        job = extract_data_points(job, div)
        job_posts.append(div['data-jk'])
        single_job_post_extension_url = "https://fr.indeed.com/viewjob?jk=" + div['data-jk']
        job['url'] = single_job_post_extension_url
        req = Request(single_job_post_extension_url, headers={'User-Agent': 'Mozilla/5.0'})
        web_page = urlopen(req).read()
        job_soup = BeautifulSoup(web_page, 'html.parser')

        for inside_div in job_soup.findAll('div', attrs={'class': 'jobsearch-jobDescriptionText'}):
            # details=inside_div.text.strip()[:100] + "..."
            details = inside_div.text.strip()
            job['details'] = details  # re.sub(' +',' ', details.replace("\n", "\n"))
        data_collected.append(job)
    return data_collected

if __name__== "__name__":
    main()
    print('----- réussi')

main()


